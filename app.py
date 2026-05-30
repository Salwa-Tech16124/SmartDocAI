import streamlit as st
import os
import tempfile
import base64
from utils.convertapi_utils import (
    convert_docx_to_pdf_and_download,
    convert_pdf_to_docx_and_download,
    compress_pdf_with_convertapi
)
from utils.agents import document_assistant_agent, summarize_document_agent, resume_review_agent

# ================================
# 🌸 SmartDoc Assistant - Elegant Feminine Theme
# ================================
st.set_page_config(page_title="SmartDoc Assistant 🌸", layout="wide", page_icon="🌸")

st.markdown("""
    <style>
    @import url('https://fonts.googleapis.com/css2?family=Playfair+Display:ital,wght@0,400;0,600;0,700;1,400&family=Nunito:wght@300;400;600;700&display=swap');

    /* Color Palette Variables */
    :root {
        --bg-color: #FAF4F6;
        --card-color: #FFFDFD;
        --rose-gold: #D8A7B1;
        --soft-pink: #F5C2D4;
        --lavender: #E8DDF7;
        --text-color: #5D4B63;
    }

    /* Main Background and Typography */
    .stApp, [data-testid="stAppViewContainer"], [data-testid="stHeader"] {
        background: var(--bg-color) !important;
        background-color: var(--bg-color) !important;
        font-family: 'Nunito', sans-serif;
        color: var(--text-color);
        overflow-x: hidden;
    }
    
    /* Hide Default Streamlit Elements */
    #MainMenu {visibility: hidden;}
    footer {visibility: hidden;}
    header {background: transparent !important;}
    
    /* Typography Overrides */
    h1, h2, h3, h4, h5, h6 {
        font-family: 'Playfair Display', serif !important;
        color: var(--text-color) !important;
        font-weight: 600 !important;
    }
    p, span, div {
        font-family: 'Nunito', sans-serif;
        color: var(--text-color);
    }

    /* Hero Section */
    .hero-container {
        text-align: center;
        padding: 5rem 1rem 3rem 1rem;
        position: relative;
        z-index: 10;
    }
    .hero-title {
        font-size: 4rem !important;
        font-weight: 700 !important;
        color: var(--text-color) !important;
        margin-bottom: 0.5rem;
        letter-spacing: 0.5px;
    }
    .hero-title span {
        color: var(--rose-gold);
        font-style: italic;
    }
    .hero-subtitle {
        font-size: 1.3rem;
        color: #8A7B8E;
        line-height: 1.6;
        font-weight: 400;
        letter-spacing: 1px;
        text-align: center;
    }

    /* Floating Metrics Cards */
    .metrics-grid {
        display: grid;
        grid-template-columns: repeat(3, 1fr);
        gap: 2rem;
        margin: 0 auto 4rem auto;
        max-width: 900px;
        position: relative;
        z-index: 10;
    }
    .glass-card {
        background: rgba(255, 253, 253, 0.7);
        backdrop-filter: blur(20px);
        -webkit-backdrop-filter: blur(20px);
        border: 1px solid rgba(216, 167, 177, 0.3);
        border-radius: 20px;
        padding: 2.5rem 1rem;
        text-align: center;
        transition: all 0.4s cubic-bezier(0.165, 0.84, 0.44, 1);
        box-shadow: 0 10px 30px rgba(216, 167, 177, 0.08);
        height: 100%;
        display: flex;
        flex-direction: column;
        justify-content: center;
    }
    .glass-card:hover {
        transform: translateY(-8px);
        box-shadow: 0 15px 40px rgba(216, 167, 177, 0.2);
        background: rgba(255, 253, 253, 0.9);
        border: 1px solid var(--rose-gold);
    }
    .metric-value {
        font-size: 2.2rem;
        font-family: 'Playfair Display', serif;
        font-weight: 700;
        margin-bottom: 0.5rem;
        color: var(--rose-gold);
    }
    .metric-label {
        font-size: 0.9rem;
        color: var(--text-color);
        text-transform: uppercase;
        letter-spacing: 2px;
        font-weight: 600;
    }

    /* Streamlit Components Overrides */
    .stTabs [data-baseweb="tab-list"] {
        gap: 1.5rem;
        background-color: transparent;
        justify-content: center;
        border-bottom: 1px solid rgba(216, 167, 177, 0.2);
        padding-bottom: 1rem;
        margin-bottom: 2rem;
    }
    .stTabs [data-baseweb="tab"] {
        height: 50px;
        white-space: pre-wrap;
        background-color: transparent;
        color: #8A7B8E;
        font-family: 'Nunito', sans-serif;
        font-weight: 600;
        font-size: 1.1rem;
        padding: 0 15px;
        border: none;
        transition: all 0.3s ease;
    }
    .stTabs [aria-selected="true"] {
        background: transparent !important;
        color: var(--text-color) !important;
        border-bottom: 3px solid var(--rose-gold) !important;
    }
    
    div.stButton {
        display: flex;
        justify-content: center;
        width: 100%;
    }
    .stButton > button {
        background: linear-gradient(135deg, var(--rose-gold), var(--soft-pink));
        color: white;
        border: none;
        border-radius: 12px;
        padding: 0.8rem 2rem;
        font-family: 'Nunito', sans-serif;
        font-weight: 600;
        font-size: 1.05rem;
        letter-spacing: 0.5px;
        box-shadow: 0 6px 15px rgba(216, 167, 177, 0.3);
        transition: all 0.3s ease;
        width: max-content !important;
        margin: 0 auto !important;
        display: block !important;
    }
    .stButton > button:hover {
        background: linear-gradient(135deg, var(--soft-pink), var(--rose-gold));
        box-shadow: 0 8px 25px rgba(216, 167, 177, 0.5);
        transform: translateY(-2px);
    }
    .stButton > button * {
        color: white !important;
    }
    
    /* Elegant Uploader styling */
    div[data-testid="stFileUploader"] {
        background: var(--card-color);
        border: 1px dashed var(--rose-gold);
        border-radius: 16px;
        padding: 2rem;
        transition: all 0.3s ease;
        box-shadow: inset 0 2px 10px rgba(216, 167, 177, 0.05);
    }
    div[data-testid="stFileUploader"]:hover {
        border-color: var(--soft-pink);
        background: #FFFFFF;
        box-shadow: 0 8px 20px rgba(216, 167, 177, 0.1);
    }
    
    /* Pastel Text Inputs */
    .stTextInput input, .stTextArea textarea {
        background-color: var(--card-color) !important;
        border: 1px solid var(--lavender) !important;
        color: var(--text-color) !important;
        border-radius: 12px !important;
        padding: 1rem 1.2rem !important;
        font-family: 'Nunito', sans-serif !important;
        font-weight: 400 !important;
        box-shadow: inset 0 1px 5px rgba(216, 167, 177, 0.05) !important;
        transition: all 0.3s ease !important;
    }
    .stTextInput input:focus, .stTextArea textarea:focus {
        border-color: var(--rose-gold) !important;
        box-shadow: 0 0 0 2px rgba(216, 167, 177, 0.2) !important;
        background-color: #ffffff !important;
    }
    
    /* History List */
    .history-item {
        background: var(--card-color);
        border: 1px solid var(--lavender);
        border-left: 4px solid var(--rose-gold);
        padding: 15px 20px;
        margin-bottom: 12px;
        border-radius: 12px;
        color: var(--text-color);
        font-weight: 600;
        font-size: 0.95rem;
        box-shadow: 0 2px 10px rgba(216, 167, 177, 0.05);
        transition: transform 0.2s ease;
    }
    .history-item:hover {
        transform: translateX(4px);
    }

    /* Ambient Background Effects (Butterflies, Petals, Sparkles) */
    .ambient-fx {
        position: fixed;
        top: 0;
        left: 0;
        width: 100vw;
        height: 100vh;
        pointer-events: none;
        z-index: 1;
        overflow: hidden;
    }
    .petal {
        position: absolute;
        font-size: 1.5rem;
        opacity: 0.6;
        animation: fall linear infinite;
    }
    .butterfly {
        position: absolute;
        font-size: 2rem;
        opacity: 0.5;
        animation: flutter ease-in-out infinite alternate;
    }
    .sparkle {
        position: absolute;
        font-size: 1.2rem;
        opacity: 0.8;
        animation: twinkle 3s ease-in-out infinite alternate;
    }
    
    @keyframes fall {
        0% { transform: translateY(-50px) rotate(0deg) translateX(0); }
        100% { transform: translateY(100vh) rotate(360deg) translateX(100px); }
    }
    @keyframes flutter {
        0% { transform: translate(0, 0) scale(1); }
        100% { transform: translate(30px, -20px) scale(1.1); }
    }
    @keyframes twinkle {
        0% { opacity: 0.2; transform: scale(0.8); }
        100% { opacity: 1; transform: scale(1.2); }
    }
    }
    </style>
""", unsafe_allow_html=True)

# Ambient Elements Injection
st.markdown("""
    <div class="ambient-fx">
        <div class="petal" style="left: 10%; animation-duration: 12s; animation-delay: 0s;">🌸</div>
        <div class="petal" style="left: 30%; animation-duration: 15s; animation-delay: 4s;">🌸</div>
        <div class="petal" style="left: 50%; animation-duration: 11s; animation-delay: 2s;">🌸</div>
        <div class="petal" style="left: 70%; animation-duration: 14s; animation-delay: 6s;">🌸</div>
        <div class="petal" style="left: 90%; animation-duration: 13s; animation-delay: 1s;">🌸</div>
        <div class="butterfly" style="top: 15%; left: 20%; animation-duration: 4s;">🦋</div>
        <div class="butterfly" style="top: 60%; left: 80%; animation-duration: 5s;">🦋</div>
        <div class="sparkle" style="top: 25%; left: 75%; animation-delay: 0s;">✨</div>
        <div class="sparkle" style="top: 70%; left: 15%; animation-delay: 1.5s;">✨</div>
        <div class="sparkle" style="top: 45%; left: 45%; animation-delay: 0.7s;">✨</div>
    </div>
""", unsafe_allow_html=True)

# ================================
# 🔊 Success Sound Function
# ================================
def play_success_sound():
    try:
        file_path = "static/success.mp3"
        if os.path.exists(file_path):
            with open(file_path, "rb") as sound_file:
                b64 = base64.b64encode(sound_file.read()).decode()
                md = f"""
                    <audio autoplay>
                        <source src="data:audio/mp3;base64,{b64}" type="audio/mpeg">
                    </audio>
                """
                st.markdown(md, unsafe_allow_html=True)
    except Exception as e:
        pass

# ================================
# 🌸 Hero & Metrics Section
# ================================
st.markdown("""
<div class="hero-container">
    <h1 class="hero-title">SmartDoc <span>Assistant</span></h1>
    <p class="hero-subtitle">Your Elegant AI Document Companion. Gracefully convert, translate, and review your documents with a touch of magic.</p>
</div>

<div class="metrics-grid">
    <div class="glass-card">
        <div class="metric-value">✨ Infinite</div>
        <div class="metric-label">Creativity</div>
    </div>
    <div class="glass-card">
        <div class="metric-value">🌸 Instant</div>
        <div class="metric-label">Processing</div>
    </div>
    <div class="glass-card">
        <div class="metric-value">🎀 Perfect</div >
        <div class="metric-label">Precision</div>
    </div>
</div>
""", unsafe_allow_html=True)

# ================================
# 🔐 ConvertAPI Token
# ================================
CONVERTAPI_BEARER_TOKEN = "hngs4CnuwWDhA8NUsmaoGLSjp2L3ELW2"

if "conversion_history" not in st.session_state:
    st.session_state.conversion_history = []


# ================================
# 🎀 Elegant Workspaces
# ================================
st.markdown("<h3 style='text-align: center; margin-top: 2rem; margin-bottom: 2rem;'>Select an Assistant 🌸</h3>", unsafe_allow_html=True)

tab1, tab2 = st.tabs(["Document Boutique 📝", "Career Stylist 👗"])

with tab1:
    st.markdown("<p style='text-align: center; margin-bottom: 2rem;'>Describe what you need, and your AI assistant will handle the rest beautifully.</p>", unsafe_allow_html=True)
    
    col1, col2 = st.columns([1, 1], gap="large")
    with col1:
        agent_file = st.file_uploader("Upload Your Document 🌸", type=["pdf", "docx", "txt"], key="agent_doc")
    with col2:
        user_instruction = st.text_input("Your Request", placeholder="E.g., 'Summarize this beautifully' ✨", key="doc_instruction")
        st.write("") # spacing
        exec_btn = st.button("✨ Begin Magic", key="exec_magic")
    
    if exec_btn:
        if agent_file and user_instruction:
            with st.spinner("🌸 Your assistant is working..."):
                action_data = document_assistant_agent(user_instruction)
                tool = action_data.get("tool", "unknown")
                
                st.success(f"🎀 Action identified: **{tool}**")
                
                # Helper to extract text
                def extract_text(f):
                    f.seek(0)
                    if f.name.endswith(".txt"): return f.read().decode("utf-8", errors="ignore")
                    elif f.name.endswith(".docx"):
                        from docx import Document
                        doc = Document(f)
                        return "\n".join([p.text for p in doc.paragraphs])
                    elif f.name.endswith(".pdf"):
                        import fitz
                        doc = fitz.open(stream=f.read(), filetype="pdf")
                        return "\n".join([page.get_text() for page in doc])
                    return ""

                if tool == "summarize_document":
                    with st.spinner("📝 Composing Summary..."):
                        text = extract_text(agent_file)
                        st.markdown("---")
                        st.markdown("### ✨ Elegant Summary")
                        st.write(summarize_document_agent(text))
                        st.session_state.conversion_history.append(f"✨ Summarized: {agent_file.name}")
                elif tool in ["docx_to_pdf", "pdf_to_docx", "compress_pdf", "translate_document"]:
                    with tempfile.NamedTemporaryFile(delete=False, suffix="."+agent_file.name.split('.')[-1]) as tmp_file:
                        agent_file.seek(0)
                        tmp_file.write(agent_file.read())
                        tmp_path = tmp_file.name
                    os.makedirs("downloads", exist_ok=True)
                    
                    if tool == "docx_to_pdf":
                        success, result = convert_docx_to_pdf_and_download(tmp_path, "downloads", CONVERTAPI_BEARER_TOKEN)
                        if success:
                            st.success("🌸 PDF gracefully generated!")
                            st.download_button("🎀 Download Document", open(result, "rb"), file_name=os.path.basename(result))
                            st.session_state.conversion_history.append(f"🌸 DOCX ➡️ PDF: {agent_file.name}")
                        else:
                            st.error(f"❌ Operation failed: {result}")
                            
                    elif tool == "pdf_to_docx":
                        success, result = convert_pdf_to_docx_and_download(tmp_path, "downloads", CONVERTAPI_BEARER_TOKEN)
                        if success:
                            st.success("🌸 Word Document gracefully generated!")
                            st.download_button("🎀 Download Document", open(result, "rb"), file_name=os.path.basename(result))
                            st.session_state.conversion_history.append(f"🌸 PDF ➡️ DOCX: {agent_file.name}")
                        else:
                            st.error(f"❌ Operation failed: {result}")
                            
                    elif tool == "compress_pdf":
                        success, result = compress_pdf_with_convertapi(tmp_path, "downloads", CONVERTAPI_BEARER_TOKEN, "medium")
                        if success:
                            st.success("🌸 Document beautifully compressed!")
                            st.download_button("🎀 Download Document", open(result, "rb"), file_name=os.path.basename(result))
                            st.session_state.conversion_history.append(f"🌸 Compressed: {agent_file.name}")
                        else:
                            st.error(f"❌ Operation failed: {result}")
                            
                    elif tool == "translate_document":
                        target_lang = action_data.get("target_lang", "hi")
                        try:
                            from docx import Document
                            from utils.editor import translate_paragraphs, save_document
                            doc = Document(tmp_path)
                            doc = translate_paragraphs(doc, target_lang=target_lang)
                            output_path = os.path.join("downloads", "agent_translated_" + agent_file.name)
                            save_document(doc, output_path)
                            st.success(f"🌸 Elegantly translated to {target_lang.upper()}!")
                            st.download_button("🎀 Download Document", open(output_path, "rb"), file_name=os.path.basename(output_path))
                            st.session_state.conversion_history.append(f"🌸 Translated ({target_lang}): {agent_file.name}")
                        except Exception as e:
                            st.error(f"❌ Operation failed (Ensure file is DOCX): {str(e)}")
                else:
                    st.warning("⚠️ Pardon me, I didn't quite catch that. Could you rephrase? 🎀")
        else:
            st.warning("🌸 Please upload a document and provide a request.")

with tab2:
    st.markdown("<p style='text-align: center; margin-bottom: 2rem;'>Receive exquisite recruiter-grade insights to refine your career profile.</p>", unsafe_allow_html=True)
    
    c1, c2 = st.columns([1, 1], gap="large")
    with c1:
        resume_file = st.file_uploader("Upload Your Resume 🦋", type=["pdf", "docx", "txt"], key="agent_resume")
    with c2:
        jd_input = st.text_area("Target Job Description ✨", height=115, placeholder="Paste the job description for tailored advice...")
    
    st.write("")
    col_btn1, col_btn2, col_btn3 = st.columns([1,2,1])
    with col_btn2:
        if st.button("🔎 Review Profile", key="exec_review"):
            if resume_file:
                with st.spinner("🎀 Reviewing your gorgeous profile..."):
                    def extract_text(f):
                        f.seek(0)
                        if f.name.endswith(".txt"): return f.read().decode("utf-8", errors="ignore")
                        elif f.name.endswith(".docx"):
                            from docx import Document
                            doc = Document(f)
                            return "\n".join([p.text for p in doc.paragraphs])
                        elif f.name.endswith(".pdf"):
                            import fitz
                            doc = fitz.open(stream=f.read(), filetype="pdf")
                            return "\n".join([page.get_text() for page in doc])
                        return ""
                        
                    res_text = extract_text(resume_file)
                    feedback = resume_review_agent(res_text, jd_input)
                    st.markdown("---")
                    st.markdown(feedback)
                    st.balloons()
                    play_success_sound()
                    st.session_state.conversion_history.append(f"🎀 Career Profile Reviewed: {resume_file.name}")
            else:
                st.warning("🌸 Kindly upload your resume first.")

# ================================
# 🕘 Activity Logs
# ================================
st.markdown("---")
st.markdown("<h3 style='text-align: center;'>🗄️ Your Elegant Logbook</h3>", unsafe_allow_html=True)
st.write("")

if st.session_state.conversion_history:
    html_history = "<div style='max-width: 800px; margin: 0 auto;'>"
    for record in reversed(st.session_state.conversion_history):
        html_history += f"<div class='history-item'>{record}</div>"
    html_history += "</div>"
    st.markdown(html_history, unsafe_allow_html=True)
else:
    st.markdown("<p style='text-align: center; color: var(--text-color);'>Your logbook is currently empty. ✨</p>", unsafe_allow_html=True)

st.write("")
col_log1, col_log2, col_log3 = st.columns([1,1,1])
with col_log2:
    if st.button("🧹 Clear Logbook", key="clear_hist"):
        st.session_state.conversion_history.clear()
        st.success("Logbook beautifully cleared!")
        st.snow()
        play_success_sound()

# ================================
# 👩‍💻 Footer
# ================================
st.markdown("---")
st.markdown("<div style='text-align:center; color:var(--text-color); font-size:0.9rem; font-family:\"Nunito\", sans-serif;'>Crafted with elegance 🌸 · SmartDoc Assistant</div>", unsafe_allow_html=True)
